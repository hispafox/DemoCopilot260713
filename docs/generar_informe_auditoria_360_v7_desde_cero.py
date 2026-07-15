from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

COLOR_TEXTO = RGBColor(0x2F, 0x2F, 0x2F)
COLOR_TITULO = RGBColor(0x0F, 0x4C, 0x81)
COLOR_SUBTITULO = RGBColor(0x1F, 0x63, 0x9A)
COLOR_CABECERA = "0F4C81"
COLOR_BLOQUE = "EAF2FA"


def sombrear_parrafo(parrafo, color_hex: str) -> None:
    p_pr = parrafo._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


def sombrear_celda(celda, color_hex: str) -> None:
    tc_pr = celda._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


def aplicar_estilos(documento: Document) -> None:
    normal = documento.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_TEXTO
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    titulo = documento.styles["Title"]
    titulo.font.name = "Calibri"
    titulo._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    titulo.font.size = Pt(24)
    titulo.font.bold = True
    titulo.font.color.rgb = COLOR_TITULO

    h1 = documento.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_TITULO
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(4)

    h2 = documento.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_SUBTITULO
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(3)


def add_titulo(documento: Document) -> None:
    p = documento.add_paragraph("Informe de Auditoria de Requisitos 360", style="Title")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    meta = documento.add_paragraph(
        "Fecha: 2026-07-15\n"
        "Proyecto: DemoCopilot260713\n"
        "Documento base: docs/documento-requisitos-aplicacion.md\n"
        "Repositorio evaluado: hispafox/DemoCopilot260713"
    )
    meta.paragraph_format.space_after = Pt(10)

    resumen = documento.add_paragraph(
        "Lectura en 60 segundos: usted tiene una base funcional solida en backend y frontend, "
        "con una brecha principal en trazabilidad operativa de GitHub y coberturas parciales "
        "en edicion UI, pruebas de actualizacion y evidencia de RNF no funcionales."
    )
    sombrear_parrafo(resumen, COLOR_BLOQUE)


def add_resumen_ejecutivo(documento: Document) -> None:
    documento.add_paragraph("1. Resumen Ejecutivo", style="Heading 1")

    documento.add_paragraph(
        "Esta auditoria cruza requisitos, codigo e issues para medir cumplimiento real. "
        "La cobertura estricta (solo requisitos en verde) es 80.0% (28/35)."
    )

    for item in [
        "Verde (completo): 28",
        "Amarillo (parcial): 6",
        "Rojo (nulo): 0",
        "Gris (no verificable): 1",
    ]:
        p = documento.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_hallazgos(documento: Document) -> None:
    documento.add_paragraph("2. Hallazgos Priorizados", style="Heading 1")

    documento.add_paragraph("Critico", style="Heading 2")
    p = documento.add_paragraph(style="List Number")
    p.add_run(
        "Desalineacion de gestion: 24/24 issues siguen abiertos pese a evidencia parcial o completa en codigo."
    )

    documento.add_paragraph("Alto", style="Heading 2")
    for texto in [
        "HU-04/RF-04 (editar tarea): backend implementado, flujo de edicion ausente en UI.",
        "RNF-03 (pruebas): faltan escenarios explicitos PUT 200 y PUT 404.",
    ]:
        p = documento.add_paragraph(style="List Number")
        p.add_run(texto)

    documento.add_paragraph("Medio", style="Heading 2")
    for texto in [
        "RNF-07: fechas UTC/ISO correctas en backend, pero sin contexto explicito de zona horaria en UI.",
        "HU-03: existe endpoint por id, pero falta experiencia de consulta por id en UI.",
        "RNF-08: falta evidencia minima de accesibilidad con casos especificos.",
    ]:
        p = documento.add_paragraph(style="List Number")
        p.add_run(texto)


def add_trazabilidad(documento: Document) -> None:
    documento.add_paragraph("3. Matriz de Trazabilidad Esencial", style="Heading 1")

    tabla = documento.add_table(rows=1, cols=5)
    tabla.style = "Table Grid"

    encabezados = [
        "Requisito",
        "Estado codigo",
        "Estado issues",
        "Riesgo",
        "Accion recomendada",
    ]
    for i, texto in enumerate(encabezados):
        celda = tabla.rows[0].cells[i]
        celda.text = texto
        sombrear_celda(celda, COLOR_CABECERA)
        for run in celda.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True
            run.font.size = Pt(10)

    filas = [
        ("HU-01 Crear tarea", "Verde", "Verde", "Bajo", "Mantener pruebas y cerrar issue con evidencia"),
        ("HU-03 Consultar por id", "Amarillo", "Verde", "Medio", "Agregar flujo UI de detalle o documentar alcance API-only"),
        ("HU-04 Editar tarea", "Amarillo", "Verde", "Alto", "Completar edicion en frontend y pruebas de flujo"),
        ("RNF-03 Pruebas", "Amarillo", "Verde", "Medio", "Agregar PUT 200/404 y casos negativos de contrato"),
        ("RNF-04 Rendimiento", "Gris", "Amarillo", "Medio", "Definir metrica y evidencia repetible"),
        ("RNF-07 UTC + ISO", "Amarillo", "Verde", "Medio", "Mostrar zona horaria de forma explicita en UI"),
        ("API PUT /api/tareas/{id}", "Verde", "Verde", "Medio", "Agregar test de 404 explicito"),
    ]

    for requisito, est_codigo, est_issue, riesgo, accion in filas:
        row = tabla.add_row().cells
        row[0].text = requisito
        row[1].text = est_codigo
        row[2].text = est_issue
        row[3].text = riesgo
        row[4].text = accion


def add_desalineaciones(documento: Document) -> None:
    documento.add_paragraph("4. Desalineaciones Detectadas", style="Heading 1")

    for item in [
        "El tablero GitHub no refleja el avance real evidenciado en codigo y pruebas.",
        "HU-04/RF-04 tiene brecha entre backend listo y UI incompleta.",
        "RNF-04 y RNF-08 carecen de evidencia fuerte y trazabilidad operacional clara.",
        "Los issues #20 a #24 amplian alcance sobre el documento base 1.0.",
    ]:
        p = documento.add_paragraph(style="List Number")
        p.add_run(item)


def add_plan(documento: Document) -> None:
    documento.add_paragraph("5. Plan de Accion Priorizado", style="Heading 1")

    tabla = documento.add_table(rows=1, cols=4)
    tabla.style = "Table Grid"
    headers = ["Prioridad", "Accion", "Tipo", "Resultado esperado"]
    for i, texto in enumerate(headers):
        c = tabla.rows[0].cells[i]
        c.text = texto
        sombrear_celda(c, "DCEAF7")
        for run in c.paragraphs[0].runs:
            run.font.bold = True

    acciones = [
        ("Alta", "Actualizar trazabilidad operativa en GitHub", "Issue/proceso", "Estado de tablero alineado con evidencia"),
        ("Alta", "Completar HU-04/RF-04 en frontend", "Codigo + test", "Flujo de edicion usable con prueba"),
        ("Media", "Ampliar pruebas de API en PUT y 404", "Test", "Cobertura de contratos criticos"),
        ("Media", "Formalizar RNF-04 con metrica objetiva", "Doc + issue + test", "Umbral de rendimiento verificable"),
        ("Media", "Cerrar brecha RNF-07/RNF-08 en UI", "Codigo + test + issue", "Zona horaria explicita y chequeo minimo de accesibilidad"),
    ]

    for prioridad, accion, tipo, resultado in acciones:
        r = tabla.add_row().cells
        r[0].text = prioridad
        r[1].text = accion
        r[2].text = tipo
        r[3].text = resultado


def add_conclusion(documento: Document) -> None:
    documento.add_paragraph("6. Conclusion", style="Heading 1")
    cierre = documento.add_paragraph(
        "Usted ya tiene una base tecnica solida para el alcance inicial. El mayor riesgo actual no es "
        "la implementacion funcional, sino la brecha entre ejecucion real y estado operativo en GitHub, "
        "junto con huecos puntuales en edicion UI y evidencia de calidad no funcional."
    )
    sombrear_parrafo(cierre, COLOR_BLOQUE)


def generar(ruta_salida: Path) -> None:
    documento = Document()
    aplicar_estilos(documento)

    add_titulo(documento)
    add_resumen_ejecutivo(documento)
    add_hallazgos(documento)
    add_trazabilidad(documento)
    add_desalineaciones(documento)
    add_plan(documento)
    add_conclusion(documento)

    ruta_salida.parent.mkdir(parents=True, exist_ok=True)
    documento.save(str(ruta_salida))


if __name__ == "__main__":
    salida = Path("docs/informe-auditoria-requisitos-360-2026-07-15-desde-cero-v7.docx")
    generar(salida)
    print(f"OK: Documento generado en {salida}")
