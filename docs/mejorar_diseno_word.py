from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

COLOR_TEXTO_BASE = RGBColor(0x44, 0x44, 0x44)
COLOR_ACENTO = RGBColor(0xE8, 0x6A, 0x10)
COLOR_ENCABEZADO_TABLA_TEXTO = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_FILA_ALTERNA = "F7F7F7"
COLOR_BLOQUE_SUAVE = "FBE3CE"
COLOR_BLOQUE_INFO = "FFF4E8"


def establecer_sombreado_parrafo(parrafo, color_hex: str) -> None:
    p_pr = parrafo._p.get_or_add_pPr()
    sombreado = p_pr.find(qn("w:shd"))
    if sombreado is None:
        sombreado = OxmlElement("w:shd")
        p_pr.append(sombreado)
    sombreado.set(qn("w:val"), "clear")
    sombreado.set(qn("w:color"), "auto")
    sombreado.set(qn("w:fill"), color_hex)


def establecer_sombreado_celda(celda, color_hex: str) -> None:
    tc_pr = celda._tc.get_or_add_tcPr()
    sombreado = tc_pr.find(qn("w:shd"))
    if sombreado is None:
        sombreado = OxmlElement("w:shd")
        tc_pr.append(sombreado)
    sombreado.set(qn("w:val"), "clear")
    sombreado.set(qn("w:color"), "auto")
    sombreado.set(qn("w:fill"), color_hex)


def copiar_fuente(origen, destino) -> None:
    destino.name = origen.name
    destino.size = origen.size
    destino.bold = origen.bold
    destino.italic = origen.italic
    destino.underline = origen.underline
    if origen.color is not None:
        destino.color.rgb = origen.color.rgb


def copiar_formato_parrafo(origen, destino) -> None:
    propiedades = [
        "alignment",
        "left_indent",
        "right_indent",
        "first_line_indent",
        "keep_together",
        "keep_with_next",
        "page_break_before",
        "widow_control",
        "space_before",
        "space_after",
        "line_spacing",
        "line_spacing_rule",
    ]
    for propiedad in propiedades:
        setattr(destino, propiedad, getattr(origen, propiedad))


def copiar_margenes(documento_origen: Document, documento_plantilla: Document) -> None:
    seccion_plantilla = documento_plantilla.sections[0]
    for seccion in documento_origen.sections:
        seccion.top_margin = seccion_plantilla.top_margin
        seccion.bottom_margin = seccion_plantilla.bottom_margin
        seccion.left_margin = seccion_plantilla.left_margin
        seccion.right_margin = seccion_plantilla.right_margin


def sincronizar_estilos_con_plantilla(documento: Document, plantilla: Document) -> None:
    estilos_objetivo = [
        "Normal",
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "List Paragraph",
    ]

    for nombre in estilos_objetivo:
        if nombre not in documento.styles or nombre not in plantilla.styles:
            continue
        estilo_destino = documento.styles[nombre]
        estilo_origen = plantilla.styles[nombre]
        copiar_fuente(estilo_origen.font, estilo_destino.font)
        copiar_formato_parrafo(
            estilo_origen.paragraph_format, estilo_destino.paragraph_format
        )


def normalizar_estilos_base(documento: Document, plantilla: Document) -> None:
    sincronizar_estilos_con_plantilla(documento, plantilla)


def aplicar_tipografia_parrafos(documento: Document) -> None:
    for parrafo in documento.paragraphs:
        estilo = (parrafo.style.name if parrafo.style is not None else "") or ""
        es_titulo = estilo in {"Title", "Heading 1", "Heading 2", "Heading 3"}

        for run in parrafo.runs:
            if run.font.name is None and "Normal" in documento.styles:
                run.font.name = documento.styles["Normal"].font.name
            if run.font.size is None and "Normal" in documento.styles:
                run.font.size = documento.styles["Normal"].font.size

            # Evita contraste bajo en titulos si el run no tiene color definido.
            if es_titulo and run.font.color.rgb is None:
                run.font.color.rgb = COLOR_ACENTO


def aplicar_formato_tablas(documento: Document) -> None:
    for tabla in documento.tables:
        if len(tabla.rows) == 0:
            continue

        fila_encabezado = tabla.rows[0]
        for celda in fila_encabezado.cells:
            establecer_sombreado_celda(celda, "E86A10")
            for parrafo in celda.paragraphs:
                parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                for run in parrafo.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
                    run.font.bold = True
                    run.font.color.rgb = COLOR_ENCABEZADO_TABLA_TEXTO
                    if run.font.size is None:
                        run.font.size = Pt(10)

        for i, fila in enumerate(tabla.rows[1:], start=1):
            if i % 2 == 1:
                color = COLOR_FILA_ALTERNA
            else:
                color = "FFFFFF"

            for celda in fila.cells:
                establecer_sombreado_celda(celda, color)
                for parrafo in celda.paragraphs:
                    for run in parrafo.runs:
                        run.font.name = "Calibri"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
                        run.font.color.rgb = COLOR_TEXTO_BASE
                        if run.font.size is None:
                            run.font.size = Pt(10)


def aplicar_encabezado_pie(documento: Document) -> None:
    for seccion in documento.sections:
        for contenedor in [seccion.header, seccion.footer]:
            for parrafo in contenedor.paragraphs:
                for run in parrafo.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
                    run.font.size = Pt(9)
                    run.font.color.rgb = COLOR_TEXTO_BASE


def es_estilo_titulo(parrafo) -> bool:
    estilo = (parrafo.style.name if parrafo.style is not None else "") or ""
    return estilo in {"Heading 1", "Heading 2", "Heading 3"}


def resaltar_secciones_clave(documento: Document) -> None:
    palabras_clave = [
        "resumen ejecutivo",
        "hallazgos por severidad",
        "plan de accion",
        "conclusion",
    ]

    for indice, parrafo in enumerate(documento.paragraphs):
        texto = parrafo.text.strip().lower()
        if not texto:
            continue

        if any(clave in texto for clave in palabras_clave):
            establecer_sombreado_parrafo(parrafo, COLOR_BLOQUE_SUAVE)
            for run in parrafo.runs:
                run.font.bold = True
                run.font.color.rgb = COLOR_ACENTO


def insertar_resumen_acciones(documento: Document) -> None:
    titulo = documento.add_paragraph("Tablero de prioridades")
    if "Heading 2" in documento.styles:
        titulo.style = documento.styles["Heading 2"]

    tabla = documento.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    cabecera = tabla.rows[0].cells
    cabecera[0].text = "Foco"
    cabecera[1].text = "Prioridad"
    cabecera[2].text = "Impacto esperado"

    filas = [
        ("Trazabilidad operativa GitHub", "Alta", "Alinea estado real del avance"),
        ("HU-04 / RF-04 en frontend", "Alta", "Cierra brecha funcional visible"),
        ("Pruebas PUT y 404", "Media", "Reduce riesgo en contratos criticos"),
    ]

    for foco, prioridad, impacto in filas:
        fila = tabla.add_row().cells
        fila[0].text = foco
        fila[1].text = prioridad
        fila[2].text = impacto


def mejorar_documento(
    ruta_origen: Path,
    ruta_plantilla: Path,
    ruta_salida: Path,
    resaltar_claves: bool = False,
    agregar_paneles: bool = False,
) -> None:
    if not ruta_origen.exists():
        raise FileNotFoundError(
            f"No existe el .docx origen: {ruta_origen}. "
            "Segun la regla, solo entonces se podria convertir desde .md."
        )

    if not ruta_plantilla.exists():
        raise FileNotFoundError(f"No existe la plantilla: {ruta_plantilla}")

    documento = Document(str(ruta_origen))
    plantilla = Document(str(ruta_plantilla))

    copiar_margenes(documento, plantilla)
    normalizar_estilos_base(documento, plantilla)
    aplicar_tipografia_parrafos(documento)
    aplicar_formato_tablas(documento)
    aplicar_encabezado_pie(documento)

    if resaltar_claves:
        resaltar_secciones_clave(documento)

    if agregar_paneles:
        insertar_resumen_acciones(documento)

    ruta_salida.parent.mkdir(parents=True, exist_ok=True)
    documento.save(str(ruta_salida))


def construir_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Mejora el diseno de un .docx manteniendo su contenido semantico."
    )
    parser.add_argument("--origen", required=True, help="Ruta del .docx origen")
    parser.add_argument("--plantilla", required=True, help="Ruta de la plantilla .docx")
    parser.add_argument("--salida", required=True, help="Ruta de salida del .docx mejorado")
    parser.add_argument(
        "--resaltar-claves",
        action="store_true",
        help="Resalta secciones clave sin modificar contenido.",
    )
    parser.add_argument(
        "--agregar-paneles",
        action="store_true",
        help="Agrega tabla adicional de prioridades al final.",
    )
    return parser


def main() -> int:
    parser = construir_parser()
    args = parser.parse_args()

    ruta_origen = Path(args.origen)
    ruta_plantilla = Path(args.plantilla)
    ruta_salida = Path(args.salida)

    try:
        mejorar_documento(
            ruta_origen,
            ruta_plantilla,
            ruta_salida,
            resaltar_claves=args.resaltar_claves,
            agregar_paneles=args.agregar_paneles,
        )
    except Exception as ex:
        print(f"ERROR: {ex}")
        return 1

    print(f"OK: Documento mejorado generado en {ruta_salida}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
